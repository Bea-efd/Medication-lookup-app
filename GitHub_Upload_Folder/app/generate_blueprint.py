import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def main():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Medication Lookup App: Logic & Blueprint', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Intro
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "This document serves as an easy-to-understand blueprint of the logic powering the Medication Lookup App. "
        "The application is designed to be a secure, offline-capable tool that accurately cross-references medication "
        "names and doses against a set of curated, trusted documents (such as NHS spreadsheets, text files, and BNF web links)."
    )
    doc.add_paragraph(
        "The core philosophy of the app is Strict Accuracy. It will not guess answers. If it cannot find "
        "the exact match for a medication and its specific dose inside the permitted data sources, it returns 'Not Found'."
    )
    
    # User Interface
    doc.add_heading('2. What the User Sees (The Interface)', level=1)
    doc.add_paragraph(
        "The application is divided into three distinct sections (Tabs) to serve different tasks:"
    )
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Single Lookup Tab: ').bold = True
    p.add_run('This is for quick questions. A user types in a single medication name (e.g., "Amoxicillin") and a dose '
              '(e.g., "500mg"). The app instantly queries all active sources and displays the formulation, price, pack size, '
              'and ATC category.')
              
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Bulk Processing Tab: ').bold = True
    p.add_run('Designed for heavy workloads. A user can upload an Excel spreadsheet containing hundreds of medication queries. '
              'The app will automatically process every single row, find the matching data, and output a completed Excel sheet '
              'with all the missing answers filled in.')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Source Library Tab (Admin Mode): ').bold = True
    p.add_run('The control center for data. This is where a manager can upload new pricing PDFs, NHS Excel sheets, or '
              'BNF web links. Each source can be toggled "Active" or "Inactive", completely controlling what knowledge '
              'the app has access to.')

    # The Engine
    doc.add_heading('3. Under the Hood (The Brain / Search Logic)', level=1)
    doc.add_paragraph(
        "When a user asks the app to look up a medication, a specific chain of thought takes place:"
    )
    
    doc.add_heading('Step A: Reading the Texts (The KnowledgeBase)', level=2)
    doc.add_paragraph(
        "Before searching, the app converts any active Excel file, PDF, or text file into a flat, readable 'Text Corpus'. "
        "This ensures that the app treats all data equally, regardless of the original file format."
    )
    
    doc.add_heading('Step B: Smart Web Scraping (BNF Live Mode)', level=2)
    doc.add_paragraph(
        "If the user has enabled internet-based searching, and a BNF source is active, the app intelligently acts "
        "like a human navigating the BNF website:"
    )
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('It reconstructs the drug name into a web address and goes straight to the medicinal-forms page.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('It ignores all generic information and "Active Ingredients" sections to avoid confusing data.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('It binds the Price (£) and Pack Size specifically to the full formulation title of the drug, ensuring '
              'prices are never mismatched to the wrong variant of the drug.')
              
    doc.add_heading('Step C: The Strict Matcher (The Logic)', level=2)
    doc.add_paragraph(
        "With all the data prepared, the app runs its strict matching logic:"
    )
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Name Verification: ').bold = True
    p.add_run('The app scans every line of text. It checks if ALL WORDS of the requested medication are present. '
              'If you search for "Amoxicillin Capsule", a line with just "Amoxicillin" is rejected.')
              
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Dose Verification: ').bold = True
    p.add_run('It heavily validates the specific dose. It checks for variations (e.g., "500 mg", "500mg"). '
              'Critically, it only checks the name of the product for the dose, NOT the surrounding generic data, '
              'eliminating false matches.')
              
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Data Extraction: ').bold = True
    p.add_run('Once a perfect row is found, it uses pattern recognition to rip out the details. It spots currency '
              'symbols (£/€/$) for Price. It spots keywords like "tablets", "capsules", or "ml" for Package Size. '
              'It spots terms like "gastro-resistant" or "dispersible" to identify the Formulation.')

    doc.add_heading('4. Architecture Flowchart (Summary)', level=1)
    doc.add_paragraph(
        "1. User Inputs Data (Single or Bulk) --> 2. App gathers active data from the Database --> "
        "3. App pulls static data from internal storage & dynamically scrapes BNF (if allowed) --> "
        "4. Strict Keyword & Pattern Logic finds the exact line of text --> 5. Final Details (Price/Size) are shipped back to the User."
    )
    
    # Save the document
    doc.save('Medication_App_Blueprint.docx')

if __name__ == '__main__':
    main()
